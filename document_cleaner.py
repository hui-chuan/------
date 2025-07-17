#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
对排文档换行符清理工具
支持藏汉对排、汉英对排、藏汉英对排等文档的换行符清理
"""

import os
import re
import sys
import argparse
from pathlib import Path
from typing import List, Tuple, Optional

# 导入docx处理库（需要安装：pip install python-docx）
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：未安装python-docx库，将无法处理docx文件")
    print("请运行：pip install python-docx")

def clean_docx_file(file_path: str) -> str:
    """读取docx文件内容"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx库未安装，无法处理docx文件")
    
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

def save_docx_file(content: str, file_path: str) -> None:
    """保存内容到docx文件"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx库未安装，无法处理docx文件")
    
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

class LanguageDetector:
    """语言检测器，支持藏文、汉文、英文的检测"""
    
    def __init__(self):
        # 藏文Unicode范围
        self.tibetan_ranges = [
            (0x0F00, 0x0FFF),  # 藏文基本范围
            (0x1000, 0x109F),  # 缅甸文（可能包含藏文扩展）
        ]
        
        # 汉文Unicode范围
        self.chinese_ranges = [
            (0x4E00, 0x9FFF),  # CJK统一汉字
            (0x3400, 0x4DBF),  # CJK扩展A
            (0x20000, 0x2A6DF), # CJK扩展B
            (0x2A700, 0x2B73F), # CJK扩展C
            (0x2B740, 0x2B81F), # CJK扩展D
            (0x2B820, 0x2CEAF), # CJK扩展E
            (0x3000, 0x303F),  # CJK符号和标点
            (0xFF00, 0xFFEF),  # 全角字符
        ]
        
        # 英文Unicode范围
        self.english_ranges = [
            (0x0041, 0x005A),  # 大写字母A-Z
            (0x0061, 0x007A),  # 小写字母a-z
            (0x0020, 0x002F),  # 空格和标点
            (0x0030, 0x0039),  # 数字0-9
            (0x003A, 0x0040),  # 标点符号
            (0x005B, 0x0060),  # 标点符号
            (0x007B, 0x007E),  # 标点符号
        ]
        
        # 日期格式正则表达式
        self.date_patterns = [
            r'^\d{4}年\d{1,2}月\d{1,2}日$',  # 2019年4月11日
            r'^\d{4}年\d{1,2}月$',           # 2019年4月
            r'^\d{1,2}月\d{1,2}日$',         # 4月11日
            r'^[A-Z][a-z]+\s+\d{1,2},\s+\d{4}$',  # April 11, 2019
            r'^[A-Z][a-z]+\s+\d{4}$',       # April 2019
            r'^\d{1,2}/\d{1,2}/\d{4}$',     # 11/4/2019
            r'^\d{4}-\d{1,2}-\d{1,2}$',     # 2019-04-11
        ]
    
    def is_date_line(self, text: str) -> bool:
        """检测文本是否为日期格式"""
        text = text.strip()
        if not text:
            return False
            
        for pattern in self.date_patterns:
            if re.match(pattern, text):
                return True
        return False

    def detect_char_language(self, char: str) -> str:
        """检测单个字符的语言类型"""
        if not char or char.isspace():
            return "space"
        
        code_point = ord(char)
        
        # 检测藏文
        for start, end in self.tibetan_ranges:
            if start <= code_point <= end:
                return "tibetan"
        
        # 检测汉文
        for start, end in self.chinese_ranges:
            if start <= code_point <= end:
                return "chinese"
        
        # 检测英文
        for start, end in self.english_ranges:
            if start <= code_point <= end:
                return "english"
        
        return "unknown"
    
    def is_punctuation(self, char: str) -> bool:
        """判断字符是否为标点符号"""
        if not char:
            return False
        
        code_point = ord(char)
        
        # 英文标点符号范围
        english_punct_ranges = [
            (0x0020, 0x002F),  # 空格和基本标点
            (0x003A, 0x0040),  # :;<=>?@
            (0x005B, 0x0060),  # [\]^_`
            (0x007B, 0x007E),  # {|}~
        ]
        
        # 中文标点符号范围
        chinese_punct_ranges = [
            (0x3000, 0x303F),  # CJK符号和标点
            (0xFF00, 0xFF0F),  # 全角ASCII标点前半部分
            (0xFF1A, 0xFF20),  # 全角ASCII标点中间部分
            (0xFF3B, 0xFF40),  # 全角ASCII标点后半部分
            (0xFF5B, 0xFF65),  # 全角ASCII标点结尾部分
        ]
        
        # 检查是否为标点符号
        for start, end in english_punct_ranges + chinese_punct_ranges:
            if start <= code_point <= end:
                return True
        
        return False
    
    def is_digit(self, char: str) -> bool:
        """判断字符是否为数字"""
        return char.isdigit()
    
    def detect_text_language(self, text: str, ignore_punctuation: bool = False) -> str:
        """检测文本的主要语言类型（改进版：数字和标点符号不参与权重计算）"""
        if not text or text.isspace():
            return "space"
        
        # 首先检查是否为日期
        if self.is_date_line(text):
            # 如果是中文日期格式，返回中文
            if re.search(r'[年月日]', text):
                return "chinese"
            # 否则返回特殊标记，表示这是日期行
            return "date"
        
        lang_counts = {"tibetan": 0, "chinese": 0, "english": 0, "unknown": 0}
        
        for char in text:
            if not char.isspace():
                # 完全忽略数字字符，不参与任何计算
                if self.is_digit(char):
                    continue
                
                # 完全忽略标点符号，不参与任何计算
                if self.is_punctuation(char):
                    continue
                    
                # 只有真正的语言字符才参与计算
                lang = self.detect_char_language(char)
                if lang in ["tibetan", "chinese", "english"]:
                    lang_counts[lang] += 1
                elif lang == "unknown":
                    lang_counts["unknown"] += 1
        
        # 统计总的有效语言字符数
        total_lang_chars = sum(lang_counts.values())
        
        # 如果没有任何语言字符，可能是纯数字或标点符号
        if total_lang_chars == 0:
            return "numeric"
        
        # 返回占比最高的语言
        max_lang = max(lang_counts, key=lang_counts.get)
        return max_lang if lang_counts[max_lang] > 0 else "unknown"
    
    def remove_parentheses_content(self, text: str) -> str:
        """移除文本中括号及其内容，支持中文括号（）和英文括号()"""
        import re
        # 移除中文括号（）及其内容
        text = re.sub(r'（[^（）]*）', '', text)
        # 移除英文括号()及其内容
        text = re.sub(r'\([^()]*\)', '', text)
        return text
    
    def has_mixed_languages(self, text: str) -> bool:
        """检测文本中是否包含多种语言（忽略括号内容）"""
        if not text or text.isspace():
            return False
        
        # 首先检查是否为日期行，日期行不算混合语言
        if self.is_date_line(text):
            return False
        
        # 移除括号内容后再进行语言检测
        text_without_parentheses = self.remove_parentheses_content(text)
        
        lang_counts = {"tibetan": 0, "chinese": 0, "english": 0}
        
        for char in text_without_parentheses:
            if not char.isspace() and not self.is_digit(char) and not self.is_punctuation(char):
                lang = self.detect_char_language(char)
                if lang in ["tibetan", "chinese", "english"]:
                    lang_counts[lang] += 1
        
        # 统计有多少种语言有字符
        languages_present = sum(1 for count in lang_counts.values() if count > 0)
        return languages_present > 1
    
    def find_language_boundaries(self, text: str) -> List[int]:
        """找到文本中语言交界处的位置（忽略括号内容）"""
        if not text or not self.has_mixed_languages(text):
            return []
        
        boundaries = []
        previous_lang = None
        in_parentheses = False
        paren_stack = []
        
        for i, char in enumerate(text):
            # 检查是否进入或退出括号
            if char in '（(':
                in_parentheses = True
                paren_stack.append(char)
                continue
            elif char in '）)':
                if paren_stack:
                    paren_stack.pop()
                if not paren_stack:
                    in_parentheses = False
                continue
            
            # 如果在括号内，跳过处理
            if in_parentheses:
                continue
                
            if char.isspace() or self.is_digit(char):
                continue
            
            current_lang = self.detect_char_language(char)
            
            # 只考虑主要语言类型
            if current_lang in ["tibetan", "chinese", "english"]:
                if previous_lang is not None and previous_lang != current_lang:
                    # 找到语言变化点，需要回溯到标点符号之后
                    boundary_pos = self._find_boundary_position(text, i)
                    if boundary_pos not in boundaries:
                        boundaries.append(boundary_pos)
                
                previous_lang = current_lang
            elif self.is_punctuation(char):
                # 标点符号不改变previous_lang，但记录位置以备后续使用
                continue
        
        return sorted(boundaries)
    
    def _find_boundary_position(self, text: str, lang_change_pos: int) -> int:
        """找到合适的语言边界位置，处理标点符号的情况"""
        # 向前查找最近的标点符号
        for i in range(lang_change_pos - 1, -1, -1):
            if self.is_punctuation(text[i]):
                # 找到标点符号，边界应该在标点符号之后
                return i + 1
            elif not text[i].isspace() and not self.is_digit(text[i]):
                # 遇到其他语言字符，边界就在语言变化点
                break
        
        return lang_change_pos
    
    def get_lines_around_newline(self, text: str, newline_pos: int) -> Tuple[str, str]:
        """获取换行符前后的完整行内容"""
        # 找到前一个换行符的位置
        prev_newline = text.rfind('\n', 0, newline_pos)
        if prev_newline == -1:
            prev_newline = 0
        else:
            prev_newline += 1  # 跳过换行符本身
        
        # 找到后一个换行符的位置
        next_newline = text.find('\n', newline_pos + 1)
        if next_newline == -1:
            next_newline = len(text)
        
        # 获取前一行和后一行的内容
        before_line = text[prev_newline:newline_pos].strip()
        after_line = text[newline_pos + 1:next_newline].strip()
        
        return before_line, after_line
    
    def is_special_single_char_line(self, line: str) -> bool:
        """检测是否为特殊的单字符行"""
        line_clean = line.strip()
        # 检查是否只包含特殊的单个字符
        special_chars = ["?", "？", "0"]
        return line_clean in special_chars
    
    def is_special_line_for_merging(self, line: str) -> bool:
        """检测是否为不应该与其他行合并的特殊行（包括单字符行和日期行）"""
        stripped = line.strip()
        # 检查单字符特殊行
        if self.is_special_single_char_line(line):
            return True
        # 检查日期行
        if self.is_date_line(stripped):
            return True
        return False

class AlignmentChecker:
    """对排文档语言交替顺序检查器"""
    
    def __init__(self, language_detector):
        self.language_detector = language_detector
        self.detected_pattern = None
        self.pattern_confidence = 0
    
    def is_special_line(self, line: str) -> bool:
        """检测是否为特殊行（空行、问号行、日期行等）"""
        line_clean = line.strip()
        return (
            not line_clean or  # 空行
            self.language_detector.is_special_single_char_line(line_clean) or  # 特殊字符行
            self.language_detector.is_date_line(line_clean)  # 日期行
        )
    
    def detect_language_pattern(self, lines: List[str]) -> Optional[List[str]]:
        """检测语言交替模式（新规则：基于前4行确定模式）"""
        if len(lines) < 4:  # 至少需要4行来检测模式
            return None
        
        # 过滤掉特殊行，只分析有效的语言行
        valid_lines = []
        for i, line in enumerate(lines):
            if not self.is_special_line(line):
                lang = self.language_detector.detect_text_language(line.strip(), ignore_punctuation=True)
                if lang != "unknown":
                    valid_lines.append((i, lang))
        
        if len(valid_lines) < 4:  # 有效行太少
            return None
        
        # 提取前4行的语言
        first_four_languages = [lang for _, lang in valid_lines[:4]]
        
        # 统计前4行中有多少种不同的语言
        unique_languages = list(set(first_four_languages))
        num_languages = len(unique_languages)
        
        detected_pattern = None
        
        if num_languages == 2:
            # 双语交替：按照前2行的交替模式作为全文交替模式
            detected_pattern = first_four_languages[:2]
            self.pattern_confidence = 95
            
        elif num_languages == 3:
            # 检查第4行是否等于第1行
            if first_four_languages[3] == first_four_languages[0]:
                # 三语交替：按照前3行的交替模式作为全文交替模式
                detected_pattern = first_four_languages[:3]
                self.pattern_confidence = 95
            else:
                # 第4行不等于第1行，无法确定明确的三语交替模式
                return None
        else:
            # 语言种类不是2或3，无法处理
            return None
        
        # 验证检测到的模式在整个文档中的匹配率
        if detected_pattern:
            all_languages = [lang for _, lang in valid_lines]
            pattern_length = len(detected_pattern)
            matches = 0
            total_checks = len(all_languages)
            
            for i, actual_lang in enumerate(all_languages):
                expected_lang = detected_pattern[i % pattern_length]
                if actual_lang == expected_lang:
                    matches += 1
            
            match_rate = matches / total_checks if total_checks > 0 else 0
            
            # 如果整体匹配率太低，调整置信度
            if match_rate < 0.7:
                self.pattern_confidence = int(match_rate * 100)
            
            return detected_pattern
        
        return None
    
    def check_alignment(self, text: str) -> str:
        """检查文档的对排一致性并标记错误"""
        lines = text.split('\n')
        
        if len(lines) < 6:
            return text  # 文档太短，无需检查
        
        # 检测语言模式
        pattern = self.detect_language_pattern(lines)
        if not pattern:
            # 如果无法检测到模式，返回原文
            return text
        
        self.detected_pattern = pattern
        pattern_length = len(pattern)
        
        # 开始检查每一行是否符合模式
        result_lines = []
        valid_line_count = 0  # 有效行计数器
        
        for i, line in enumerate(lines):
            if self.is_special_line(line):
                # 特殊行直接保留
                result_lines.append(line)
            else:
                # 检查语言是否符合模式
                current_lang = self.language_detector.detect_text_language(line.strip(), ignore_punctuation=True)
                
                if current_lang != "unknown":
                    expected_lang = pattern[valid_line_count % pattern_length]
                    
                    if current_lang == expected_lang:
                        # 语言匹配，正常添加
                        result_lines.append(line)
                    else:
                        # 语言不匹配，添加错误标记
                        result_lines.append(line + " {aligned_error}")
                    
                    valid_line_count += 1
                else:
                    # 无法识别语言的行，保留但不计入模式
                    result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def get_pattern_info(self) -> str:
        """获取检测到的模式信息"""
        if self.detected_pattern:
            pattern_str = " → ".join(self.detected_pattern)
            pattern_type = "双语交替" if len(self.detected_pattern) == 2 else "三语交替"
            return f"检测到的语言模式: {pattern_str} ({pattern_type}, 置信度: {self.pattern_confidence}%)"
        else:
            return "未检测到明确的语言交替模式"

class DocumentCleaner:
    """文档清理器"""
    
    def __init__(self):
        self.language_detector = LanguageDetector()
        self.alignment_checker = AlignmentChecker(self.language_detector)
    
    def clean_consecutive_newlines(self, text: str) -> str:
        """清理连续的换行符，只保留一个"""
        # 使用正则表达式替换连续的换行符
        cleaned = re.sub(r'\n{2,}', '\n', text)
        return cleaned
    
    def split_mixed_language_lines(self, text: str) -> str:
        """处理包含多种语言的行，在语言交界处添加换行符"""
        lines = text.split('\n')
        result_lines = []
        
        for line in lines:
            if self.language_detector.has_mixed_languages(line):
                # 找到语言边界位置
                boundaries = self.language_detector.find_language_boundaries(line)
                
                if boundaries:
                    # 按照边界位置分割行
                    split_parts = []
                    start = 0
                    
                    for boundary in boundaries:
                        if start < boundary:
                            split_parts.append(line[start:boundary].strip())
                        start = boundary
                    
                    # 添加最后一部分
                    if start < len(line):
                        split_parts.append(line[start:].strip())
                    
                    # 过滤掉空白部分，然后添加到结果中
                    valid_parts = [part for part in split_parts if part.strip()]
                    result_lines.extend(valid_parts)
                else:
                    # 没有找到有效边界，保持原行
                    result_lines.append(line)
            else:
                # 单一语言的行，直接添加
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def should_remove_newline(self, text: str, newline_pos: int) -> bool:
        """判断是否应该删除指定位置的换行符（新逻辑：只处理单一语言的行）"""
        before_line, after_line = self.language_detector.get_lines_around_newline(text, newline_pos)
        
        # 检查是否为空行
        if not before_line or not after_line:
            return False
        
        # 检查是否为特殊行（包括单字符行和日期行），如果是则跳过处理
        if (self.language_detector.is_special_line_for_merging(before_line) or 
            self.language_detector.is_special_line_for_merging(after_line)):
            return False
        
        # 检查前后行是否包含多种语言，如果任一行包含多种语言，则不删除换行
        if (self.language_detector.has_mixed_languages(before_line) or 
            self.language_detector.has_mixed_languages(after_line)):
            return False
        
        # 检测语言类型，忽略标点符号
        before_lang = self.language_detector.detect_text_language(before_line, ignore_punctuation=True)
        after_lang = self.language_detector.detect_text_language(after_line, ignore_punctuation=True)
        
        # 如果前后都是同一种语言且都是单一语言，则删除换行符
        if before_lang == after_lang and before_lang != "unknown":
            return True
        
        return False
    
    def clean_same_language_newlines(self, text: str) -> str:
        """清理同种语言内的换行符"""
        result = []
        i = 0
        
        while i < len(text):
            if text[i] == '\n':
                if self.should_remove_newline(text, i):
                    # 删除换行符，但保留一个空格以防止单词连接
                    result.append(' ')
                else:
                    result.append('\n')
            else:
                result.append(text[i])
            i += 1
        
        return ''.join(result)
    
    def clean_text(self, text: str) -> str:
        """清理文本中的多余换行符（新逻辑）"""
        # 步骤1：清理连续的换行符
        step1 = self.clean_consecutive_newlines(text)
        
        # 步骤2：处理包含多种语言的行，在语言交界处添加换行
        step2 = self.split_mixed_language_lines(step1)
        
        # 步骤3：清理同种语言内的换行符（只对单一语言的行生效）
        step3 = self.clean_same_language_newlines(step2)
        
        # 步骤4：再次清理可能产生的连续换行符
        step4 = self.clean_consecutive_newlines(step3)
        
        # 步骤5：清理多余的空格
        step5 = re.sub(r' +', ' ', step4)  # 多个空格替换为单个空格
        step5 = re.sub(r' *\n *', '\n', step5)  # 换行符前后的空格
        
        return step5.strip()
    
    def check_alignment(self, text: str) -> Tuple[str, str]:
        """检查对排文档的语言交替一致性"""
        checked_text = self.alignment_checker.check_alignment(text)
        pattern_info = self.alignment_checker.get_pattern_info()
        return checked_text, pattern_info
    
    def clean_and_check_text(self, text: str, check_alignment: bool = True) -> Tuple[str, str]:
        """清理文本并检查对排一致性（修正逻辑：先处理后检查）"""
        if check_alignment:
            # 先进行文本清理，不带任何错误标记
            cleaned_text = self.clean_text(text)
            
            # 对清理后的最终结果进行模式检查
            final_checked_text, pattern_info = self.check_alignment(cleaned_text)
            
            return final_checked_text, pattern_info
        else:
            # 只进行基本清理
            cleaned_text = self.clean_text(text)
            return cleaned_text, ""
    
    def process_txt_file(self, file_path: str, check_alignment: bool = True) -> None:
        """处理txt文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            print(f"处理文件: {file_path}")
            if check_alignment:
                cleaned_content, pattern_info = self.clean_and_check_text(content, check_alignment=True)
                if pattern_info:
                    print(f"  {pattern_info}")
                else:
                    print("  未检测到明确的语言交替模式")
            else:
                cleaned_content = self.clean_text(content)
                print("  跳过对排检查")
            
            # 保存清理后的文件
            output_path = file_path.replace('.txt', '_cleaned.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(cleaned_content)
            
            print(f"已完成: {file_path} -> {output_path}")
            
        except Exception as e:
            print(f"处理txt文件时出错 {file_path}: {e}")
    
    def process_docx_file(self, file_path: str, check_alignment: bool = True) -> None:
        """处理docx文件"""
        if not DOCX_AVAILABLE:
            print(f"跳过docx文件 {file_path}：未安装python-docx库")
            return
        
        try:
            doc = Document(file_path)
            
            # 提取所有文本内容进行对排检查
            all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if check_alignment and all_text.strip():
                checked_text, pattern_info = self.check_alignment(all_text)
                print(f"处理文件: {file_path}")
                print(f"  {pattern_info}")
                
                # 将检查后的文本重新分配给段落
                checked_lines = checked_text.split('\n')
                paragraph_index = 0
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip() and paragraph_index < len(checked_lines):
                        # 清理文本但不再检查对排（已经检查过了）
                        cleaned_line = self.clean_text(checked_lines[paragraph_index])
                        paragraph.text = cleaned_line
                        paragraph_index += 1
            else:
                # 只进行基本清理
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        cleaned_text = self.clean_text(paragraph.text)
                        paragraph.text = cleaned_text
            
            # 保存清理后的文件
            output_path = file_path.replace('.docx', '_cleaned.docx')
            doc.save(output_path)
            
            print(f"已完成: {file_path} -> {output_path}")
            
        except Exception as e:
            print(f"处理docx文件时出错 {file_path}: {e}")
    
    def process_file(self, file_path: str, check_alignment: bool = True) -> None:
        """处理单个文件"""
        if file_path.lower().endswith('.txt'):
            self.process_txt_file(file_path, check_alignment)
        elif file_path.lower().endswith('.docx'):
            self.process_docx_file(file_path, check_alignment)
        else:
            print(f"不支持的文件格式: {file_path}")
    
    def process_directory(self, directory: str, check_alignment: bool = True) -> None:
        """处理目录中的所有文档"""
        path = Path(directory)
        
        if not path.exists():
            print(f"目录不存在: {directory}")
            return
        
        # 查找所有txt和docx文件
        txt_files = list(path.glob('*.txt'))
        docx_files = list(path.glob('*.docx'))
        
        all_files = txt_files + docx_files
        
        if not all_files:
            print(f"在目录 {directory} 中未找到txt或docx文件")
            return
        
        print(f"找到 {len(all_files)} 个文件待处理")
        
        for file_path in all_files:
            self.process_file(str(file_path), check_alignment)

def main():
    parser = argparse.ArgumentParser(description='清理对排文档中的多余换行符并检查语言交替一致性')
    parser.add_argument('path', help='要处理的文件或目录路径')
    parser.add_argument('--test', action='store_true', help='测试模式，显示清理效果但不保存文件')
    parser.add_argument('--no-alignment-check', action='store_true', help='跳过对排一致性检查')
    
    args = parser.parse_args()
    
    cleaner = DocumentCleaner()
    check_alignment = not args.no_alignment_check
    
    if os.path.isfile(args.path):
        cleaner.process_file(args.path, check_alignment)
    elif os.path.isdir(args.path):
        cleaner.process_directory(args.path, check_alignment)
    else:
        print(f"路径不存在: {args.path}")

if __name__ == "__main__":
    main() 